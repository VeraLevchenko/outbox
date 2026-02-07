from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import StreamingResponse, FileResponse
from sqlalchemy.orm import Session
from datetime import date, datetime
from pathlib import Path
from pydantic import BaseModel
import io
import uuid
import base64

from app.models.database import SessionLocal
from app.schemas.outbox_schemas import RegisterRequest, RegisterResponse
from app.services.kaiten_service import kaiten_service
from app.services.file_service import file_service
from app.services.docx_service import docx_service
from app.services.config_service import config_service
from app.services.pdf_service import pdf_service
from app.services.cryptopro_service import cryptopro_service
from app.api.auth import get_current_user
from app.api.journal import get_next_outgoing_number

router = APIRouter(prefix="/api/outbox", tags=["outbox"])

# Директория для временных зарегистрированных файлов
TEMP_FILES_DIR = Path(__file__).parent.parent.parent / "temp_files"
TEMP_FILES_DIR.mkdir(exist_ok=True)


def get_db():
    """Dependency для получения сессии БД"""
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


@router.post("/prepare-registration", response_model=RegisterResponse)
async def prepare_registration(
    request: RegisterRequest,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """
    Подготовить регистрацию документа:
    - Получить карточку и извлечь title (для поля "Кому")
    - Получить исполнителя из карточки Kaiten (для генерации номера)
    - Сгенерировать номер
    - Заменить плейсхолдеры в DOCX
    - Вернуть информацию для предпросмотра

    Args:
        request: Данные запроса (card_id)
        db: Сессия БД
        current_user: Текущий пользователь

    Returns:
        Данные регистрации с номером и датой
    """
    try:
        # 1. Получаем карточку для извлечения title
        card = await kaiten_service.get_card_by_id(request.card_id)
        if not card:
            raise HTTPException(status_code=404, detail=f"Card {request.card_id} not found")

        # Извлекаем title карточки - это поле "Кому"
        to_whom = card.get('title', '')

        # 2. Получаем исполнителя из карточки Kaiten (для генерации номера)
        executor_data = await kaiten_service.get_executor_from_card(request.card_id)

        if not executor_data:
            raise HTTPException(
                status_code=404,
                detail="Executor not found in card. Please ensure card has a member with type=2"
            )

        executor_id = executor_data.get('user_id')
        executor_name = executor_data.get('full_name')

        # 3. Генерируем следующий номер
        # Используем существующий endpoint get_next_outgoing_number
        from app.api.journal import get_next_outgoing_number as get_next_number_func
        from sqlalchemy import func
        from app.models.outbox_journal import OutboxJournal

        # Получаем правила нумерации для исполнителя
        numbering_rule = config_service.get_numbering_rule_for_executor(executor_id)

        # Извлекаем параметры из правила
        executor_code = numbering_rule.get('executor_code', '00')
        number_format = numbering_rule.get('format', '{number}-{executor_code}')
        start_number = numbering_rule.get('start_number', 1)
        reset_yearly = numbering_rule.get('reset_yearly', False)

        # Определяем фильтр для поиска максимального номера
        query = db.query(func.max(OutboxJournal.outgoing_no))

        # Если нумерация сбрасывается ежегодно, фильтруем по текущему году
        if reset_yearly:
            current_year = date.today().year
            query = query.filter(
                func.extract('year', OutboxJournal.outgoing_date) == current_year
            )

        # Сквозная нумерация независимо от исполнителя - НЕ фильтруем по executor

        # Получаем максимальный номер
        max_no = query.scalar()
        next_number = (max_no or (start_number - 1)) + 1

        # Форматируем номер согласно правилу (например, 42-10)
        formatted_number = number_format.format(
            number=next_number,
            executor_code=executor_code
        )

        # 4. Получаем текущую дату в формате ДД.ММ.ГГГГ
        today = date.today()
        outgoing_date = docx_service.format_date(today)

        # 5. Проверяем, что выбранный файл - DOCX
        if not request.selected_file_name.lower().endswith('.docx'):
            raise HTTPException(
                status_code=400,
                detail=f"Выбранный файл '{request.selected_file_name}' не является DOCX документом. Регистрировать можно только DOCX файлы с полями для заполнения."
            )

        # 6. Находим выбранный файл в карточке
        card_files = card.get('files', [])
        selected_file = None
        for file_info in card_files:
            if file_info.get('name') == request.selected_file_name:
                selected_file = file_info
                break

        if not selected_file:
            raise HTTPException(
                status_code=404,
                detail=f"Файл '{request.selected_file_name}' не найден в карточке"
            )

        # 8. Скачиваем DOCX (в mock режиме используем mock данные)
        if file_service.use_mock:
            # В mock режиме создаем простой DOCX с плейсхолдерами
            print(f"[Mock] Creating mock DOCX with placeholders for file: {request.selected_file_name}")
            docx_bytes = _create_mock_docx()
        else:
            # Скачиваем реальный файл из Kaiten
            docx_url = selected_file.get('url') or selected_file.get('path')
            if not docx_url:
                raise HTTPException(
                    status_code=404,
                    detail=f"URL файла '{request.selected_file_name}' не найден"
                )
            docx_bytes = await docx_service.download_docx_from_url(docx_url)

        # 9. Проверяем наличие плейсхолдеров
        has_placeholders = docx_service.check_has_placeholders(docx_bytes)
        if not has_placeholders:
            raise HTTPException(
                status_code=400,
                detail=f"Файл '{request.selected_file_name}' не содержит полей для заполнения ({{{{outgoing_no}}}}, {{{{outgoing_date}}}}, {{{{stamp}}}}). Регистрировать можно только шаблоны с полями."
            )

        # 10. Заменяем плейсхолдеры (пока без данных сертификата)
        modified_docx = docx_service.replace_placeholders(
            docx_bytes,
            formatted_number,
            outgoing_date,
            certificate_data=None  # Сначала создаем без подписи
        )

        # 11. Конвертируем DOCX в PDF
        print(f"[Outbox] Converting DOCX to PDF...")
        try:
            pdf_bytes = pdf_service.convert_docx_to_pdf(modified_docx)
            print(f"[Outbox] PDF created: {len(pdf_bytes)} bytes")
        except Exception as e:
            print(f"[Outbox] PDF conversion error: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Ошибка конвертации в PDF: {str(e)}"
            )

        # 12. НЕ подписываем на сервере - подпись будет создана на клиенте через браузер
        # Вместо этого просто используем DOCX без штампа ЭЦП
        print(f"[Outbox] PDF ready for client-side signing")
        modified_docx_with_stamp = modified_docx  # Используем DOCX без штампа

        # 14. Сохраняем файлы во временное хранилище
        # Убеждаемся, что директория существует
        TEMP_FILES_DIR.mkdir(exist_ok=True, parents=True)

        file_id = str(uuid.uuid4())
        # Формируем имя файла: номер_дата_оригинальное_имя
        safe_number = formatted_number.replace('/', '_').replace('\\', '_').replace('-', '_')
        safe_date = outgoing_date.replace('.', '_')

        # Санитизируем оригинальное имя файла - убираем проблемные символы
        base_name = request.selected_file_name.rsplit('.', 1)[0]  # без расширения
        # Заменяем пробелы, скобки и другие проблемные символы
        safe_base_name = base_name.replace(' ', '_').replace('(', '').replace(')', '').replace('[', '').replace(']', '')

        # Сохраняем DOCX
        docx_filename = f"{safe_number}_{safe_date}_{safe_base_name}.docx"
        docx_file_path = TEMP_FILES_DIR / f"{file_id}_{docx_filename}"
        with open(docx_file_path, 'wb') as f:
            f.write(modified_docx_with_stamp)

        # Сохраняем PDF (без подписи - будет подписан на клиенте)
        pdf_filename = f"{safe_number}_{safe_date}_{safe_base_name}.pdf"
        pdf_file_path = TEMP_FILES_DIR / f"{file_id}_{pdf_filename}"
        with open(pdf_file_path, 'wb') as f:
            f.write(pdf_bytes)

        # Подпись (.sig) НЕ создаём здесь - будет создана на клиенте через браузер

        print(f"[Outbox] Files saved:")
        print(f"  - DOCX: {docx_file_path}")
        print(f"  - PDF: {pdf_file_path}")
        print(f"  - SIG: будет создана на клиенте")

        # URL для скачивания PDF
        download_url = f"/api/outbox/download/{file_id}_{pdf_filename}"

        # URL для подписания (редирект на sign.html)
        sign_url = f"/sign.html?fileId={file_id}&pdfFile={file_id}_{pdf_filename}"

        return RegisterResponse(
            outgoing_no=next_number,
            formatted_number=formatted_number,
            outgoing_date=outgoing_date,
            executor=executor_name,
            executor_id=executor_id,
            docx_preview_url=download_url,
            sign_url=sign_url,
            file_id=file_id,
            message=f"Документ готов к подписанию. Номер: {formatted_number} от {outgoing_date}"
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f"Error preparing registration: {e}")
        raise HTTPException(status_code=500, detail=f"Error preparing registration: {str(e)}")


def _create_mock_docx() -> bytes:
    """Создать mock DOCX файл с плейсхолдерами для тестирования"""
    from docx import Document

    doc = Document()
    doc.add_heading('Исходящее письмо', 0)

    doc.add_paragraph(f'Исх. № {{{{outgoing_no}}}} от {{{{outgoing_date}}}}')
    doc.add_paragraph('')
    doc.add_paragraph('Уважаемые коллеги,')
    doc.add_paragraph('')
    doc.add_paragraph('Направляем Вам информацию по запросу.')
    doc.add_paragraph('')
    doc.add_paragraph('С уважением,')
    doc.add_paragraph('{{stamp}}')

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


class ClientSignatureUpload(BaseModel):
    """Схема для приёма подписи, созданной на клиенте"""
    file_id: str  # ID временного файла
    signature: str  # Base64 подпись
    thumbprint: str  # Отпечаток сертификата
    cn: str  # Common Name владельца сертификата
    # Дополнительные данные для записи в журнал
    card_id: int  # ID карточки Kaiten
    outgoing_no: int  # Номер документа (число)
    formatted_number: str  # Форматированный номер (например, "178-01")
    outgoing_date: str  # Дата в формате ДД.ММ.ГГГГ
    to_whom: str  # Кому (из названия карточки)
    executor: str  # Исполнитель


@router.post("/upload-client-signature")
async def upload_client_signature(
    data: ClientSignatureUpload,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """
    Принять подпись, созданную на клиенте (через браузер с КриптоПро)
    и создать запись в журнале исходящих документов

    Args:
        data: Данные подписи
        db: Сессия БД
        current_user: Текущий пользователь

    Returns:
        Результат сохранения подписи и создания записи в журнале
    """
    try:
        # Декодируем подпись из Base64
        try:
            sig_bytes = base64.b64decode(data.signature)
        except Exception as e:
            raise HTTPException(
                status_code=400,
                detail=f"Неверный формат подписи Base64: {str(e)}"
            )

        # Проверяем, что файл существует
        # Ищем файлы с этим file_id
        matching_files = list(TEMP_FILES_DIR.glob(f"{data.file_id}_*.pdf"))
        if not matching_files:
            raise HTTPException(
                status_code=404,
                detail=f"PDF файл с ID {data.file_id} не найден"
            )

        pdf_file_path = matching_files[0]

        # Создаём .sig файл рядом с PDF
        sig_file_path = pdf_file_path.with_suffix('.pdf.sig')
        sig_file_path.write_bytes(sig_bytes)

        # Логируем информацию о подписи
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        log_entry = f"""
=== Подпись получена от клиента ===
Время: {timestamp}
Владелец сертификата: {data.cn}
Отпечаток: {data.thumbprint}
PDF файл: {pdf_file_path.name}
Размер PDF: {pdf_file_path.stat().st_size} байт
Размер подписи: {len(sig_bytes)} байт
========================
"""
        print(log_entry)

        # Сохраняем в лог-файл
        log_path = TEMP_FILES_DIR / "signatures.log"
        try:
            with open(log_path, "a", encoding="utf-8") as f:
                f.write(log_entry + "\n")
        except Exception as e:
            print(f"[Outbox] Warning: Could not write to log file: {e}")

        # ========== СОЗДАНИЕ ЗАПИСИ В ЖУРНАЛЕ ==========

        # 1. Получаем карточку Kaiten для извлечения краткого содержания
        print(f"[Outbox] Getting Kaiten card {data.card_id} for journal entry...")
        card = await kaiten_service.get_card_by_id(data.card_id)
        if not card:
            raise HTTPException(status_code=404, detail=f"Card {data.card_id} not found")

        # Извлекаем краткое содержание из свойства карточки (properties)
        from app.core.config import settings as app_settings
        properties = card.get('properties', {})
        content = properties.get(app_settings.KAITEN_PROPERTY_CONTENT, '') or ''
        # Если в свойствах нет, пробуем description
        if not content:
            content = card.get('description', '') or ''

        # Формируем ссылку на карточку Kaiten
        kaiten_url = f"https://outbox.kaiten.ru/space/397084/card/{data.card_id}"

        # 2. Читаем файлы
        print(f"[Outbox] Reading PDF and SIG files...")
        pdf_bytes = pdf_file_path.read_bytes()

        # 3. Получаем приложения (все файлы из карточки, кроме основного DOCX)
        print(f"[Outbox] Getting attachments from Kaiten...")
        attachments_bytes = None
        card_files = card.get('files', [])

        # Фильтруем файлы: исключаем основной DOCX, который был зарегистрирован
        # Это определяется по имени файла из registrationResult
        # Но нам нужно знать, какой файл был выбран - это сложно определить здесь
        # Поэтому возьмём все файлы, кроме .docx файлов (так как основной уже в PDF)
        attachment_files = [f for f in card_files if not f.get('name', '').lower().endswith('.docx')]

        if attachment_files:
            print(f"[Outbox] Found {len(attachment_files)} attachments")
            # Скачиваем и упаковываем в ZIP архив
            import zipfile
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for file_info in attachment_files:
                    file_name = file_info.get('name', 'unknown')
                    file_url = file_info.get('url') or file_info.get('path')
                    if file_url:
                        try:
                            # Скачиваем файл
                            file_bytes = await file_service.download_file(file_url)
                            zip_file.writestr(file_name, file_bytes)
                            print(f"  - Added: {file_name}")
                        except Exception as e:
                            print(f"  - Failed to download {file_name}: {e}")

            attachments_bytes = zip_buffer.getvalue()
            print(f"[Outbox] Attachments archive size: {len(attachments_bytes)} bytes")

        # ========== СОХРАНЕНИЕ ФАЙЛОВ В ПАПКУ ИСХОДЯЩИХ ==========

        # 4. Создаём папку и сохраняем файлы в /mnt/doc/Исходящие/{номер}
        from app.core.config import settings
        outgoing_folder = Path(settings.OUTGOING_FILES_PATH) / data.formatted_number

        print(f"[Outbox] Creating folder: {outgoing_folder}")
        try:
            outgoing_folder.mkdir(parents=True, exist_ok=True)
        except PermissionError as e:
            error_msg = f"Нет прав на создание папки {outgoing_folder}. Создайте папку вручную и настройте права: sudo mkdir -p {Path(settings.OUTGOING_FILES_PATH)} && sudo chown -R $USER:$USER {Path(settings.OUTGOING_FILES_PATH)}"
            print(f"[Outbox] ERROR: Permission denied creating folder: {outgoing_folder}")
            print(f"[Outbox] Please run: sudo mkdir -p {Path(settings.OUTGOING_FILES_PATH)} && sudo chown -R $USER:$USER {Path(settings.OUTGOING_FILES_PATH)}")
            raise HTTPException(status_code=500, detail=error_msg)

        # Сохраняем PDF
        pdf_filename = pdf_file_path.name.replace(f"{data.file_id}_", "")  # Убираем file_id из имени
        pdf_save_path = outgoing_folder / pdf_filename
        pdf_save_path.write_bytes(pdf_bytes)
        print(f"[Outbox] Saved PDF: {pdf_save_path}")

        # Сохраняем SIG
        sig_filename = pdf_filename.replace('.pdf', '.pdf.sig')
        sig_save_path = outgoing_folder / sig_filename
        sig_save_path.write_bytes(sig_bytes)
        print(f"[Outbox] Saved SIG: {sig_save_path}")

        # Сохраняем приложения (отдельные файлы, а не архив)
        if attachment_files:
            print(f"[Outbox] Saving {len(attachment_files)} attachments to folder...")
            for file_info in attachment_files:
                file_name = file_info.get('name', 'unknown')
                file_url = file_info.get('url') or file_info.get('path')
                if file_url:
                    try:
                        # Скачиваем файл (повторно, если нужно)
                        file_bytes = await file_service.download_file(file_url)
                        attachment_save_path = outgoing_folder / file_name
                        attachment_save_path.write_bytes(file_bytes)
                        print(f"  - Saved: {file_name}")
                    except Exception as e:
                        print(f"  - Failed to save {file_name}: {e}")

        print(f"[Outbox] All files saved to: {outgoing_folder}")

        # 6. Конвертируем дату из ДД.ММ.ГГГГ в date
        from datetime import datetime as dt
        outgoing_date_obj = dt.strptime(data.outgoing_date, "%d.%m.%Y").date()

        # 7. Создаём запись в журнале
        from app.models.outbox_journal import OutboxJournal

        print(f"[Outbox] Creating journal entry...")
        journal_entry = OutboxJournal(
            outgoing_no=data.outgoing_no,  # Числовая часть (например, 178)
            formatted_number=data.formatted_number,  # Полный форматированный номер (например, "178-01")
            outgoing_date=outgoing_date_obj,
            to_whom=data.to_whom,
            executor=data.executor,
            content=content[:500] if content else None,  # Ограничиваем длину
            kaiten_card_url=kaiten_url,
            file_blob=pdf_bytes,
            sig_blob=sig_bytes,
            attachments_blob=attachments_bytes,
            folder_path=str(outgoing_folder)  # Путь к папке с файлами
        )

        db.add(journal_entry)
        db.commit()
        db.refresh(journal_entry)

        print(f"[Outbox] Journal entry created: ID={journal_entry.id}")

        return {
            "success": True,
            "message": "Подпись сохранена, файлы записаны в папку и запись добавлена в журнал",
            "pdf_file": pdf_file_path.name,
            "sig_file": sig_file_path.name,
            "folder_path": str(outgoing_folder),
            "journal_entry_id": journal_entry.id,
            "timestamp": timestamp,
            "certificate": {
                "cn": data.cn,
                "thumbprint": data.thumbprint
            }
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"[Outbox] Error uploading client signature: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Ошибка сохранения подписи: {str(e)}"
        )


@router.get("/download/{filename}")
async def download_file(filename: str):
    """
    Скачать файл из временного хранилища

    Args:
        filename: Имя файла для скачивания

    Returns:
        Файл для скачивания
    """
    # Защита от path traversal
    if ".." in filename or "/" in filename or "\\" in filename:
        raise HTTPException(status_code=400, detail="Неверное имя файла")

    file_path = TEMP_FILES_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Файл не найден")

    # Определяем MIME тип
    if filename.endswith('.pdf'):
        media_type = "application/pdf"
    elif filename.endswith('.sig'):
        media_type = "application/octet-stream"
    elif filename.endswith('.docx'):
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        media_type = "application/octet-stream"

    return FileResponse(
        file_path,
        media_type=media_type,
        filename=filename
    )


@router.get("/status/{file_id}")
async def get_file_status(file_id: str):
    """
    Получить статус файлов (PDF, DOCX, SIG) по file_id

    Args:
        file_id: ID файла

    Returns:
        Информация о существующих файлах
    """
    pdf_files = list(TEMP_FILES_DIR.glob(f"{file_id}_*.pdf"))
    docx_files = list(TEMP_FILES_DIR.glob(f"{file_id}_*.docx"))
    sig_files = list(TEMP_FILES_DIR.glob(f"{file_id}_*.pdf.sig"))

    result = {
        "file_id": file_id,
        "pdf_exists": len(pdf_files) > 0,
        "docx_exists": len(docx_files) > 0,
        "sig_exists": len(sig_files) > 0,
    }

    if pdf_files:
        pdf_path = pdf_files[0]
        result["pdf_file"] = pdf_path.name
        result["pdf_size"] = pdf_path.stat().st_size

    if docx_files:
        docx_path = docx_files[0]
        result["docx_file"] = docx_path.name
        result["docx_size"] = docx_path.stat().st_size

    if sig_files:
        sig_path = sig_files[0]
        result["sig_file"] = sig_path.name
        result["sig_size"] = sig_path.stat().st_size

    return result
