import os
import io
import httpx
from pathlib import Path
from datetime import date
from typing import Optional, Dict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.core.config import settings


class DocxService:
    """Сервис для работы с DOCX документами"""

    def __init__(self):
        self.static_path = Path(__file__).parent.parent / "static"
        self.stamp_image_path = self.static_path / "stamp.png"

    async def download_docx_from_url(self, url: str) -> bytes:
        """
        Скачать DOCX файл по URL

        Args:
            url: URL файла

        Returns:
            Содержимое файла в виде байтов
        """
        async with httpx.AsyncClient() as client:
            response = await client.get(url, timeout=30.0)
            response.raise_for_status()
            return response.content

    def check_has_placeholders(self, docx_bytes: bytes) -> bool:
        """
        Проверить, есть ли в документе плейсхолдеры для заполнения

        Args:
            docx_bytes: Содержимое DOCX файла

        Returns:
            True если найдены плейсхолдеры, False если нет
        """
        try:
            doc = Document(io.BytesIO(docx_bytes))

            # Проверяем плейсхолдеры в параграфах
            for paragraph in doc.paragraphs:
                if '{{outgoing_no}}' in paragraph.text or \
                   '{{outgoing_date}}' in paragraph.text or \
                   '{{stamp}}' in paragraph.text:
                    return True

            # Проверяем плейсхолдеры в таблицах
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if '{{outgoing_no}}' in paragraph.text or \
                               '{{outgoing_date}}' in paragraph.text or \
                               '{{stamp}}' in paragraph.text:
                                return True

            return False
        except Exception as e:
            print(f"Error checking placeholders: {e}")
            return False

    def replace_placeholders(
        self,
        docx_bytes: bytes,
        outgoing_no: str,
        outgoing_date: str,
        certificate_data: Optional[Dict] = None
    ) -> bytes:
        """
        Заменить плейсхолдеры в DOCX документе

        Args:
            docx_bytes: Содержимое DOCX файла
            outgoing_no: Исходящий номер (например, "42-10")
            outgoing_date: Дата в формате ДД.ММ.ГГГГ (например, "20.01.2026")
            certificate_data: Данные сертификата для визуализации ЭЦП

        Returns:
            Измененный DOCX файл в виде байтов
        """
        # Загружаем документ из байтов
        doc = Document(io.BytesIO(docx_bytes))

        # Заменяем текстовые плейсхолдеры в параграфах
        for paragraph in doc.paragraphs:
            if '{{outgoing_no}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{outgoing_no}}', outgoing_no)
            if '{{outgoing_date}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{outgoing_date}}', outgoing_date)

            # Заменяем {{stamp}} на изображение визуализации ЭЦП
            if '{{stamp}}' in paragraph.text:
                # Очищаем текст параграфа
                paragraph.text = ''
                # Вставляем визуализацию ЭЦП
                self._insert_stamp_visualization(paragraph, certificate_data)

        # Заменяем плейсхолдеры в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '{{outgoing_no}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{outgoing_no}}', outgoing_no)
                        if '{{outgoing_date}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{outgoing_date}}', outgoing_date)

                        # Заменяем {{stamp}} в таблицах
                        if '{{stamp}}' in paragraph.text:
                            paragraph.text = ''
                            self._insert_stamp_visualization(paragraph, certificate_data)

        # Сохраняем измененный документ в байты
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()

    def _insert_stamp_visualization(self, paragraph, certificate_data: Optional[Dict] = None):
        """
        Вставить визуализацию электронной подписи в параграф

        Args:
            paragraph: Параграф документа
            certificate_data: Данные сертификата
        """
        # Если есть готовое изображение штампа, используем его
        if self.stamp_image_path.exists():
            try:
                run = paragraph.add_run()
                run.add_picture(str(self.stamp_image_path), width=Inches(2.5))
                return
            except Exception as e:
                print(f"Error adding stamp image: {e}")

        # Если изображения нет, создаем текстовую визуализацию в рамке
        # Используем данные из certificate_data или mock данные
        cert_serial = certificate_data.get('serial', '5C6BE147FA657D807EF3A907DFB53553') if certificate_data else '5C6BE147FA657D807EF3A907DFB53553'
        cert_owner = certificate_data.get('owner', 'Левченко Вера Сергеевна') if certificate_data else 'Левченко Вера Сергеевна'
        cert_valid_from = certificate_data.get('valid_from', '16.07.2025') if certificate_data else '16.07.2025'
        cert_valid_to = certificate_data.get('valid_to', '09.10.2026') if certificate_data else '09.10.2026'

        # Добавляем текстовую визуализацию
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Заголовок
        run = paragraph.add_run('ДОКУМЕНТ ПОДПИСАН ЭЛЕКТРОННОЙ ПОДПИСЬЮ\n')
        run.bold = True
        run.font.size = Pt(10)

        # Данные сертификата
        run = paragraph.add_run(f'Сертификат {cert_serial}\n')
        run.font.size = Pt(8)

        run = paragraph.add_run(f'Владелец {cert_owner}\n')
        run.font.size = Pt(8)

        run = paragraph.add_run(f'Действителен с {cert_valid_from} по {cert_valid_to}')
        run.font.size = Pt(8)

    def format_date(self, date_obj: date) -> str:
        """
        Форматировать дату в формат ДД.ММ.ГГГГ

        Args:
            date_obj: Объект даты

        Returns:
            Строка с датой в формате ДД.ММ.ГГГГ
        """
        return date_obj.strftime('%d.%m.%Y')


# Singleton instance
docx_service = DocxService()
